"""DOCX text extraction via python-docx."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from app.parsers._versions import package_version
from app.parsers.base import DocumentParser, ParserResult
from app.services.text_extraction_service import normalize_extension


class DocxParser(DocumentParser):
    PARSER_NAME = "docx_parser"

    def __init__(self) -> None:
        self._version = package_version("python-docx")

    def supports(self, extension: str, mime_type: str | None) -> bool:
        return normalize_extension(extension) == "docx"

    def _table_to_lines(self, table) -> list[str]:  # type: ignore[no-untyped-def]
        lines: list[str] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                t = (cell.text or "").replace("\n", " ").replace("\t", " ").strip()
                cells.append(t)
            lines.append(" | ".join(cells))
        return lines

    def parse_bytes(
        self, content: bytes, filename: str, extension: str
    ) -> ParserResult:
        try:
            doc = Document(BytesIO(content))
        except Exception:
            return ParserResult(
                success=False,
                extracted_text=None,
                text_length=0,
                parser_name=self.PARSER_NAME,
                parser_version=self._version,
                metadata={},
                error_code="PARSING_FAILED",
                error_message="Document parsing failed",
            )

        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        table_lines: list[str] = []
        for ti, table in enumerate(doc.tables, start=1):
            table_lines.append(f"\n[Table {ti}]\n")
            table_lines.extend(self._table_to_lines(table))

        body = "\n".join(parts)
        if table_lines:
            body = (body + "\n" + "\n".join(table_lines)).strip()

        return ParserResult(
            success=True,
            extracted_text=body,
            text_length=len(body),
            parser_name=self.PARSER_NAME,
            parser_version=self._version,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )


__all__ = ["DocxParser"]
